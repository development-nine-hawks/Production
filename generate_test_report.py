from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

PHOTOS_DIR = r"C:\AIB\Projects\NineHawks\With Pattern\Akshat Testing Photos"
OUTPUT_PATH = r"C:\AIB\Projects\NineHawks\With Pattern\Codebase\Phase1_Test_Report_v2.docx"

test_cases = [
    {
        "id": "TC-01",
        "title": "Genuine label, photo straight-on",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "82.3%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.216",
        "color": "0.406",
        "markers": "4",
        "alignment": "perspective (0\u00b0)",
        "image": "Genuine label, photo straight-on.jpeg",
    },
    {
        "id": "TC-02",
        "title": "Genuine label, slight angle (~30\u00b0)",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "86.1%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.455",
        "color": "0.430",
        "markers": "4",
        "alignment": "perspective (0\u00b0)",
        "image": "Genuine label, slight angle (~30\u00b0).jpeg",
    },
    {
        "id": "TC-03",
        "title": "Genuine label, rotated 90\u00b0",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "83.7%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.315",
        "color": "0.398",
        "markers": "4",
        "alignment": "perspective (90\u00b0CW)",
        "image": "Genuine label, rotated 90\u00b0.jpeg",
    },
    {
        "id": "TC-04",
        "title": "Genuine label, rotated 180\u00b0",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "83.9%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.308",
        "color": "0.430",
        "markers": "4",
        "alignment": "perspective (180\u00b0)",
        "image": "Genuine label, rotated 180\u00b0.jpeg",
    },
    {
        "id": "TC-05",
        "title": "Genuine label, low light",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "78.5%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.146",
        "color": "0.126",
        "markers": "4",
        "alignment": "perspective (0\u00b0)",
        "image": "Genuine label, low light.jpeg",
    },
    {
        "id": "TC-06",
        "title": "Genuine label, glare/flash on pattern",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "80.3%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.227",
        "color": "0.193",
        "markers": "4",
        "alignment": "perspective (0\u00b0)",
        "image": "Genuine label, glare flash on pattern.jpeg",
    },
    {
        "id": "TC-07",
        "title": "Genuine label, far away (pattern small in frame)",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "84.8%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.404",
        "color": "0.378",
        "markers": "4",
        "alignment": "perspective (0\u00b0)",
        "image": "Genuine label, far away (pattern small in frame).jpeg",
    },
    {
        "id": "TC-08",
        "title": "Genuine label, motion blur",
        "expected": "AUTHENTIC",
        "verdict": "AUTHENTIC",
        "pass_fail": "PASS",
        "confidence": "81.5%",
        "moire": "1.000",
        "correlation": "1.000",
        "gradient": "0.203",
        "color": "0.346",
        "markers": "4",
        "alignment": "resize (0)",
        "image": "Genuine label, motion blur.jpeg",
    },
    {
        "id": "TC-09",
        "title": "Counterfeit label, straight-on (print \u2192 photo \u2192 reprint \u2192 photo)",
        "expected": "COUNTERFEIT",
        "verdict": "COUNTERFEIT",
        "pass_fail": "PASS",
        "confidence": "44.1%",
        "moire": "0.481",
        "correlation": "1.000",
        "gradient": "0.014",
        "color": "0.266",
        "markers": "4",
        "alignment": "resize (0)",
        "image": "Counterfeit label, straight-on (print \u2192 photo \u2192 reprint \u2192 photo).jpeg",
    },
    {
        "id": "TC-10",
        "title": "Counterfeit label, at an angle",
        "expected": "COUNTERFEIT",
        "verdict": "COUNTERFEIT",
        "pass_fail": "PASS",
        "confidence": "43.9%",
        "moire": "0.504",
        "correlation": "0.718",
        "gradient": "0.089",
        "color": "0.262",
        "markers": "4",
        "alignment": "perspective (0\u00b0)",
        "image": "Counterfeit label, at an angle.jpeg",
    },
    {
        "id": "TC-11",
        "title": "Wrong pattern ID against genuine print",
        "expected": "COUNTERFEIT",
        "verdict": "COUNTERFEIT",
        "pass_fail": "PASS",
        "confidence": "3.3%",
        "moire": "0.004",
        "correlation": "0.000",
        "gradient": "0.040",
        "color": "0.245",
        "markers": "4",
        "alignment": "crop_resize",
        "image": "Genuine label, photo straight-on.jpeg",
    },
    {
        "id": "TC-12",
        "title": "Screenshot of label on screen, photographed",
        "expected": "COUNTERFEIT",
        "verdict": "COUNTERFEIT",
        "pass_fail": "PASS",
        "confidence": "5.8%",
        "moire": "0.041",
        "correlation": "0.000",
        "gradient": "0.104",
        "color": "0.155",
        "markers": "4",
        "alignment": "resize (90CW)",
        "image": "Screenshot of label on screen, photographed.jpeg",
    },
]

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_score_row(table, label, value):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Title
title = doc.add_heading('NineHawks CDP — Phase 1 Test Report', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Print Size: 7.5mm  |  Pattern Type: Label-embedded  |  Block Size: 16×16')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

# Summary table
doc.add_heading('Summary', level=1)
summary_table = doc.add_table(rows=1, cols=5)
summary_table.style = 'Table Grid'
hdr = summary_table.rows[0].cells
for cell, text in zip(hdr, ['TC', 'Description', 'Expected', 'Actual', 'Result']):
    cell.text = text
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(cell, 'D9D9D9')

for tc in test_cases:
    row = summary_table.add_row().cells
    row[0].text = tc['id']
    row[1].text = tc['title']
    row[2].text = tc['expected']
    row[3].text = tc['verdict']
    row[4].text = tc['pass_fail']
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    if tc['pass_fail'] == 'PASS':
        set_cell_bg(row[4], 'C6EFCE')
        row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x27, 0x6F, 0x2D)
    else:
        set_cell_bg(row[4], 'FFC7CE')
        row[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x9C, 0x00, 0x06)

passed = sum(1 for tc in test_cases if tc['pass_fail'] == 'PASS')
doc.add_paragraph()
summary_para = doc.add_paragraph()
summary_run = summary_para.add_run(f'Result: {passed}/{len(test_cases)} passed')
summary_run.font.bold = True
summary_run.font.size = Pt(11)

doc.add_page_break()

# Individual test cases
doc.add_heading('Test Case Details', level=1)

for tc in test_cases:
    doc.add_heading(f"{tc['id']}: {tc['title']}", level=2)

    # Pass/Fail badge line
    pf_para = doc.add_paragraph()
    pf_run = pf_para.add_run(f"  {tc['pass_fail']}  ")
    pf_run.font.bold = True
    pf_run.font.size = Pt(10)
    if tc['pass_fail'] == 'PASS':
        pf_run.font.color.rgb = RGBColor(0x27, 0x6F, 0x2D)
    else:
        pf_run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
    exp_run = pf_para.add_run(f"   Expected: {tc['expected']}   →   Got: {tc['verdict']}   |   Confidence: {tc['confidence']}")
    exp_run.font.size = Pt(10)

    # Two-column layout: image left, scores right
    layout = doc.add_table(rows=1, cols=2)
    layout.style = 'Table Grid'

    # Left: image
    img_cell = layout.rows[0].cells[0]
    img_path = os.path.join(PHOTOS_DIR, tc['image'])
    if os.path.exists(img_path):
        img_para = img_cell.paragraphs[0]
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(img_path, width=Inches(2.8))
    else:
        img_cell.text = f"[Image not found: {tc['image']}]"

    # Right: scores table
    score_cell = layout.rows[0].cells[1]
    score_cell.paragraphs[0].clear()
    scores_table = score_cell.add_table(rows=0, cols=2)
    scores_table.style = 'Table Grid'

    rows_data = [
        ('Verdict', tc['verdict']),
        ('Confidence', tc['confidence']),
        ('Moire (65%)', tc['moire']),
        ('Correlation (10%)', tc['correlation']),
        ('Gradient (15%)', tc['gradient']),
        ('Color (10%)', tc['color']),
        ('Markers Found', tc['markers']),
        ('Alignment', tc['alignment']),
    ]
    for label, value in rows_data:
        r = scores_table.add_row().cells
        r[0].text = label
        r[1].text = value
        r[0].paragraphs[0].runs[0].font.size = Pt(9)
        r[1].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(r[0], 'F2F2F2')

    doc.add_paragraph()

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
